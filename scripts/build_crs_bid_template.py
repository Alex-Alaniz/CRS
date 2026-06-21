from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "CRS-Bid-Proposal-Template.docx"

NAVY = RGBColor(8, 16, 32)
BLUE = RGBColor(46, 116, 181)
BLUE_DARK = RGBColor(31, 77, 120)
GOLD = RGBColor(212, 169, 85)
RED = RGBColor(196, 30, 58)
MUTED = RGBColor(91, 107, 128)
LIGHT_FILL = "F4F6F9"
WHITE = RGBColor(240, 244, 248)


def set_run(run, size=None, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_text(cell, text, bold=False, color=None, size=10.5, fill=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, color=color or NAVY, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "C9D2DF")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            set_cell_margins(cell)
    set_table_borders(table)


def add_para(doc, text="", style=None, align=None, before=None, after=None, line_spacing=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run(run, size=11, color=NAVY)
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    if before is not None:
        fmt.space_before = Pt(before)
    if after is not None:
        fmt.space_after = Pt(after)
    if line_spacing is not None:
        fmt.line_spacing = line_spacing
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        run = p.add_run(item)
        set_run(run, size=11, color=NAVY)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=BLUE_DARK, fill=LIGHT_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, BLUE_DARK, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    run = header.add_run("Controlled Risk Services LLC | Bid Proposal Template")
    set_run(run, size=9, color=MUTED, bold=True)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    run = footer.add_run("Integrity. Protection. Excellence.")
    set_run(run, size=9, color=MUTED)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_cover(doc):
    p = add_para(doc, "Controlled Risk Services LLC", align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    set_run(p.runs[0], size=12, color=MUTED, bold=True)

    title = add_para(doc, "Bid Proposal Template", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    set_run(title.runs[0], size=26, color=NAVY, bold=True)

    sub = add_para(doc, "Safety Staffing, Field Support, and Risk Management Services", align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    set_run(sub.runs[0], size=14, color=MUTED)

    tag = add_para(doc, "Prepared for client, agency, and project-specific bid responses", align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    set_run(tag.runs[0], size=10.5, color=GOLD, bold=True)

    table = doc.add_table(rows=4, cols=2)
    set_table_width(table, [1.35, 4.95])
    rows = [
        ("Client / Agency", "[CLIENT / AGENCY NAME]"),
        ("Project / Solicitation", "[SOLICITATION ID / PROJECT NAME]"),
        ("Primary Contact", "Carlos Castro\ncarlos.castro@controlledriskservices.com\n346-436-3004"),
        ("Submission Date", "[MONTH DAY, YEAR]"),
    ]
    for r, (label, value) in enumerate(rows):
        set_cell_text(table.rows[r].cells[0], label, bold=True, color=BLUE_DARK, fill=LIGHT_FILL)
        set_cell_text(table.rows[r].cells[1], value)

    note = add_para(doc, "Use this template as a controlled starting point. Replace bracketed fields, confirm pricing assumptions, attach required resumes/certifications, and review every compliance response before submission.", after=12)
    set_run(note.runs[0], size=10.5, color=MUTED, italic=True)


def build():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    add_cover(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc, "Controlled Risk Services LLC proposes to provide safety leadership, field safety professionals, and project-risk support for [CLIENT / PROJECT]. CRS brings construction and industrial field experience, practical safety governance, and staffing coverage designed to protect people while keeping the project moving.")
    add_bullets(doc, [
        "Protect workers, client assets, and project continuity.",
        "Provide qualified safety professionals who integrate with the site team.",
        "Document risks, actions, and outcomes clearly enough for owners, contractors, and compliance stakeholders.",
    ])

    add_heading(doc, "2. Project Understanding", 1)
    add_heading(doc, "Client Need", 2)
    add_para(doc, "[Summarize the client's stated need, project type, location, schedule, and critical safety or staffing gaps.]")
    add_heading(doc, "CRS Response", 2)
    add_para(doc, "CRS will support the project through safety management, field coverage, compliance documentation, and practical coordination with the client team.")

    add_heading(doc, "3. Scope of Services", 1)
    add_table(
        doc,
        ["Service Area", "CRS Responsibility", "Deliverable"],
        [
            ("Safety consultant / SME", "Subject-matter support for hazard identification, compliance questions, and field controls.", "Safety recommendations, field notes, action logs."),
            ("Field safety professional", "On-site safety presence, observations, daily communication, and workforce support.", "Daily reports, inspection notes, toolbox-talk support."),
            ("Construction / project support", "Coordination, schedule awareness, contractor communication, and issue escalation.", "Meeting notes, risk register updates, escalation summaries."),
            ("Cost / staffing support", "Track staffing levels, expected hours, and pricing assumptions.", "Staffing plan and pricing schedule."),
        ],
        [1.65, 2.75, 2.2],
    )

    add_heading(doc, "4. Staffing Plan", 1)
    add_table(
        doc,
        ["Role", "Qty", "Est. Hours", "Key Qualifications", "Notes"],
        [
            ("Safety Consultant SME", "[ ]", "[ ]", "OSHA/compliance, incident review, program support", "[ ]"),
            ("Field Safety Professional", "[ ]", "[ ]", "Site observations, JHA/JSA support, daily reporting", "[ ]"),
            ("Construction Manager", "[ ]", "[ ]", "Contractor coordination, schedule support, field leadership", "[ ]"),
            ("Cost Manager", "[ ]", "[ ]", "Budget tracking, staffing cost summaries", "[ ]"),
        ],
        [1.55, 0.5, 0.85, 2.45, 1.15],
    )

    add_heading(doc, "5. Technical Approach", 1)
    add_heading(doc, "Mobilization", 2)
    add_para(doc, "CRS will confirm project scope, reporting expectations, site rules, communication channels, and staffing schedule before field mobilization.")
    add_heading(doc, "Field Execution", 2)
    add_para(doc, "CRS personnel will perform jobsite observations, participate in safety meetings, support pre-task planning, document risks, and escalate issues that require client action.")
    add_heading(doc, "Reporting", 2)
    add_para(doc, "CRS will provide reports at an agreed cadence. Reports may include daily summaries, safety observations, corrective actions, photos if authorized, incident notes, and open-item tracking.")

    add_heading(doc, "6. Compliance Matrix", 1)
    add_table(
        doc,
        ["Requirement", "CRS Response", "Evidence / Attachment", "Status"],
        [
            ("[RFP requirement]", "[How CRS satisfies it]", "[Resume, certification, policy, or report sample]", "[Compliant / Exception / Clarify]"),
            ("[RFP requirement]", "[How CRS satisfies it]", "[Attachment reference]", "[Status]"),
            ("[RFP requirement]", "[How CRS satisfies it]", "[Attachment reference]", "[Status]"),
        ],
        [1.65, 2.25, 2.05, 0.95],
    )

    add_heading(doc, "7. Pricing Schedule", 1)
    add_table(
        doc,
        ["Labor Category", "Rate", "Estimated Hours", "Extended Total"],
        [
            ("Safety Consultant SME", "$[ ] / hr", "[ ]", "$[ ]"),
            ("Field Safety Professional", "$[ ] / hr", "[ ]", "$[ ]"),
            ("Construction Manager", "$[ ] / hr", "[ ]", "$[ ]"),
            ("Cost Manager", "$[ ] / hr", "[ ]", "$[ ]"),
            ("Expenses / Travel", "$[ ]", "[ ]", "$[ ]"),
            ("Total Proposed Price", "", "", "$[ ]"),
        ],
        [2.4, 1.25, 1.5, 1.45],
    )

    add_heading(doc, "8. Assumptions and Exclusions", 1)
    add_bullets(doc, [
        "Pricing assumes [standard hours / shift schedule / project duration].",
        "Client will provide site access, orientation requirements, and project-specific safety procedures before mobilization.",
        "CRS pricing excludes permits, third-party testing, specialized equipment, and travel unless specifically listed in the pricing schedule.",
        "Any scope changes, schedule extensions, or additional roles will be handled through written change approval.",
    ])

    add_heading(doc, "9. Differentiators", 1)
    add_bullets(doc, [
        "Field-first safety leadership grounded in construction and industrial project experience.",
        "Flexible safety staffing model for short-term, long-term, and surge coverage needs.",
        "Practical reporting that supports both field execution and client accountability.",
        "Commitment to integrity, protection, and excellence.",
    ])

    add_heading(doc, "10. Submission Checklist", 1)
    add_bullets(doc, [
        "[ ] Confirm solicitation number and due date.",
        "[ ] Confirm required forms, certifications, and representations.",
        "[ ] Attach CRS capability statement.",
        "[ ] Attach resumes and certifications for proposed personnel.",
        "[ ] Complete compliance matrix.",
        "[ ] Review pricing math and assumptions.",
        "[ ] Confirm delivery method and submission contact.",
    ])

    add_heading(doc, "11. Acceptance", 1)
    add_table(
        doc,
        ["Authorized Party", "Name / Title", "Signature", "Date"],
        [
            ("Client Representative", "", "", ""),
            ("Controlled Risk Services", "", "", ""),
        ],
        [1.65, 2.1, 1.65, 1.2],
    )

    doc.core_properties.title = "Controlled Risk Services Bid Proposal Template"
    doc.core_properties.author = "Controlled Risk Services LLC"
    doc.core_properties.subject = "Bid proposal template"
    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
